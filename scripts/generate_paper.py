#!/usr/bin/env python3
"""
Generate a formatted Chinese academic paper .docx from a JSON input file.

Usage:
    python generate_paper.py input.json output.docx

Requires: python-docx (pip install python-docx)
"""

import json
import sys
import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE


# Font size constants (in Pt)
SAN_HAO = Pt(16)      # 三号
SI_HAO = Pt(14)       # 四号
XIAO_SI = Pt(12)      # 小四
WU_HAO = Pt(10.5)     # 五号
LIU_HAO = Pt(7.5)     # 六号

# Font names
HEI_TI = '黑体'
SONG_TI = '宋体'
TIMES = 'Times New Roman'


def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, font_name_cn, font_name_en, size, bold=False, italic=False):
    """Set font for a run with both Chinese and English font names."""
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.name = font_name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    rFonts.set(qn('w:cs'), font_name_en)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_paragraph_with_font(doc, text, font_cn, font_en, size, bold=False,
                             italic=False, alignment=None, line_spacing=1.5):
    """Add a paragraph with specified font settings."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold, italic)
    return p


def add_section_heading(doc, number, title, level):
    """Add a numbered section heading with proper formatting."""
    if level == 1:
        font_size = SI_HAO
        font_cn = HEI_TI
        full_title = f"{number} {title}"
    elif level == 2:
        font_size = XIAO_SI
        font_cn = HEI_TI
        full_title = f"{number} {title}"
    else:
        font_size = XIAO_SI
        font_cn = SONG_TI
        full_title = f"{number} {title}"

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.5, space_before=6, space_after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(full_title)
    set_run_font(run, font_cn, TIMES, font_size, bold=True)
    return p


def format_table_cell(cell, text, font_cn=SONG_TI, font_en=TIMES, size=XIAO_SI,
                       bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Format a table cell with proper font and alignment."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    set_run_font(run, font_cn, font_en, size, bold)


def apply_three_line_table(table):
    """Apply 三线表 formatting: thick top, thin header-bottom, thick bottom."""
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell,
                            top={"sz": 0, "val": "none"},
                            bottom={"sz": 0, "val": "none"},
                            left={"sz": 0, "val": "none"},
                            right={"sz": 0, "val": "none"})

    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})

    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})

    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})


def build_reference_text(ref, index):
    """Build reference text from structured data based on type."""
    t = ref.get('type', 'other')
    authors = ref.get('authors', '')
    title = ref.get('title', '')
    ref_type = ref.get('ref_type', '')

    type_codes = {
        'journal': 'J', 'book': 'M', 'conference': 'C', 'dissertation': 'D',
        'newspaper': 'N', 'electronic': 'EB/OL', 'patent': 'P',
        'standard': 'S', 'report': 'R', 'collection': 'G', 'other': 'Z'
    }
    code = ref_type or type_codes.get(t, 'Z')

    if t == 'journal':
        journal = ref.get('journal', '')
        year = ref.get('year', '')
        volume = ref.get('volume', '')
        issue = ref.get('issue', '')
        pages = ref.get('pages', '')
        issn = ref.get('issn', '')
        issn_str = f" (S{issn})" if issn else ""
        vol_issue = f"{volume}({issue})" if issue else volume
        return f"[{index}] {authors}.{title}[{code}].{journal}{issn_str},{year},{vol_issue}:{pages}."

    elif t == 'book':
        publisher = ref.get('publisher', '')
        location = ref.get('location', '')
        year = ref.get('year', '')
        pages = ref.get('pages', '')
        edition = ref.get('edition', '')
        translator = ref.get('translator', '')
        edition_str = f".{edition}" if edition else ""
        translator_str = f".{translator} 译" if translator else ""
        pages_str = f":{pages}" if pages else ""
        return f"[{index}] {authors}.{title}[{code}]{translator_str}{edition_str}.{location}:{publisher},{year}{pages_str}."

    elif t == 'conference':
        editor = ref.get('editor', '')
        source = ref.get('source', '')
        location = ref.get('location', '')
        publisher = ref.get('publisher', '')
        year = ref.get('year', '')
        pages = ref.get('pages', '')
        return f"[{index}] {authors}.{title}[{code}]//{editor}.{source}.{location}:{publisher},{year}:{pages}."

    elif t == 'dissertation':
        location = ref.get('location', '')
        institution = ref.get('institution', '')
        year = ref.get('year', '')
        return f"[{index}] {authors}.{title}[{code}].{location}:{institution},{year}."

    elif t == 'newspaper':
        newspaper = ref.get('newspaper', '')
        date = ref.get('date', '')
        edition = ref.get('edition', '')
        edition_str = f"({edition})" if edition else ""
        return f"[{index}] {authors}.{title}[{code}].{newspaper},{date}{edition_str}."

    elif t == 'electronic':
        url = ref.get('url', '')
        pub_date = ref.get('pub_date', '')
        access_date = ref.get('access_date', '')
        date_str = f"{pub_date}/{access_date}" if access_date else pub_date
        return f"[{index}] {authors}.{title}[{code}].{url},{date_str}."

    elif t == 'patent':
        country = ref.get('country', '')
        patent_no = ref.get('patent_no', '')
        date = ref.get('date', '')
        return f"[{index}] {authors}.{title}[{code}].{country},{patent_no},{date}."

    elif t == 'standard':
        std_code = ref.get('std_code', '')
        location = ref.get('location', '')
        publisher = ref.get('publisher', '')
        year = ref.get('year', '')
        return f"[{index}] {authors}.{std_code} {title}[{code}].{location}:{publisher},{year}."

    else:
        location = ref.get('location', '')
        publisher = ref.get('publisher', '')
        year = ref.get('year', '')
        return f"[{index}] {authors}.{title}[Z].{location}:{publisher},{year}."


def generate_front_matter(doc, data):
    """Generate title, abstracts, and keywords with Roman page numbering."""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    sectPr = section._sectPr
    pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="lowerRoman"/>')
    sectPr.append(pgNumType)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in [run, run2, run3]:
        set_run_font(r, SONG_TI, TIMES, LIU_HAO)

    title_cn = data.get('title_cn', '')
    p_title_cn = add_paragraph_with_font(
        doc, title_cn, HEI_TI, TIMES, SAN_HAO,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5
    )
    p_title_cn.paragraph_format.space_after = Pt(6)

    title_en = data.get('title_en', '')
    p_title_en = add_paragraph_with_font(
        doc, title_en, TIMES, TIMES, SAN_HAO,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5
    )
    p_title_en.paragraph_format.space_after = Pt(12)

    abstract_cn = data.get('abstract_cn', {})
    p_abs_label = doc.add_paragraph()
    set_paragraph_spacing(p_abs_label, 1.5)
    p_abs_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_label = p_abs_label.add_run('摘要')
    set_run_font(run_label, HEI_TI, TIMES, XIAO_SI, bold=True)
    run_sep = p_abs_label.add_run('：')
    set_run_font(run_sep, SONG_TI, TIMES, XIAO_SI)

    abs_text = abstract_cn.get('text', '')
    innovations = abstract_cn.get('innovations', [])

    p_abs_body = doc.add_paragraph()
    set_paragraph_spacing(p_abs_body, 1.5)
    p_abs_body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_body = p_abs_body.add_run(abs_text)
    set_run_font(run_body, SONG_TI, TIMES, XIAO_SI)

    if innovations:
        p_abs_body.add_run(' ').font.size = XIAO_SI
        run_inno_label = p_abs_body.add_run('创新点：')
        set_run_font(run_inno_label, HEI_TI, TIMES, XIAO_SI, bold=True)
        inno_text = '；'.join(innovations) + '。'
        run_inno = p_abs_body.add_run(inno_text)
        set_run_font(run_inno, SONG_TI, TIMES, XIAO_SI)

    keywords_cn = data.get('keywords_cn', [])
    p_kw = doc.add_paragraph()
    set_paragraph_spacing(p_kw, 1.5)
    p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_kw_label = p_kw.add_run('关键词')
    set_run_font(run_kw_label, HEI_TI, TIMES, XIAO_SI, bold=True)
    run_kw_sep = p_kw.add_run('：')
    set_run_font(run_kw_sep, SONG_TI, TIMES, XIAO_SI)
    run_kw_text = p_kw.add_run('；'.join(keywords_cn))
    set_run_font(run_kw_text, SONG_TI, TIMES, XIAO_SI)
    p_kw.paragraph_format.space_after = Pt(12)

    abstract_en = data.get('abstract_en', {})
    p_en_label = doc.add_paragraph()
    set_paragraph_spacing(p_en_label, 1.5)
    p_en_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_en = p_en_label.add_run('Abstract')
    set_run_font(run_en, TIMES, TIMES, XIAO_SI, bold=True)
    run_en_sep = p_en_label.add_run(': ')
    set_run_font(run_en_sep, TIMES, TIMES, XIAO_SI)

    abs_en_text = abstract_en.get('text', '')
    innovations_en = abstract_en.get('innovations', [])

    p_en_body = doc.add_paragraph()
    set_paragraph_spacing(p_en_body, 1.5)
    p_en_body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_en_body = p_en_body.add_run(abs_en_text)
    set_run_font(run_en_body, TIMES, TIMES, XIAO_SI)

    if innovations_en:
        p_en_body.add_run(' ').font.size = XIAO_SI
        run_en_inno_label = p_en_body.add_run('Innovations: ')
        set_run_font(run_en_inno_label, TIMES, TIMES, XIAO_SI, bold=True)
        inno_en_text = '; '.join(innovations_en) + '.'
        run_en_inno = p_en_body.add_run(inno_en_text)
        set_run_font(run_en_inno, TIMES, TIMES, XIAO_SI, italic=True)

    keywords_en = data.get('keywords_en', [])
    p_kw_en = doc.add_paragraph()
    set_paragraph_spacing(p_kw_en, 1.5)
    p_kw_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_kw_en_label = p_kw_en.add_run('Key words')
    set_run_font(run_kw_en_label, TIMES, TIMES, XIAO_SI, bold=True)
    run_kw_en_sep = p_kw_en.add_run(': ')
    set_run_font(run_kw_en_sep, TIMES, TIMES, XIAO_SI)
    run_kw_en_text = p_kw_en.add_run('; '.join(keywords_en))
    set_run_font(run_kw_en_text, TIMES, TIMES, XIAO_SI)

    doc.add_page_break()


def generate_toc(doc, data):
    """Generate table of contents page."""
    p_toc_title = add_paragraph_with_font(
        doc, '目  录', HEI_TI, TIMES, SAN_HAO,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5
    )
    p_toc_title.paragraph_format.space_after = Pt(12)

    sections = data.get('sections', [])

    def _collect_headings(children, prefix='', indent_level=0):
        entries = []
        counter = [0, 0, 0]
        for item in children:
            if item.get('type') == 'section' or 'level' in item:
                level = item.get('level', 1)
                if level == 1:
                    counter[0] += 1
                    counter[1] = 0
                    counter[2] = 0
                    num = str(counter[0])
                elif level == 2:
                    counter[1] += 1
                    counter[2] = 0
                    num = f"{counter[0]}.{counter[1]}"
                else:
                    counter[2] += 1
                    num = f"{counter[0]}.{counter[1]}.{counter[2]}"

                entries.append({
                    'number': num,
                    'title': item.get('title', ''),
                    'level': level
                })

                sub = item.get('children', [])
                if sub:
                    sub_entries = _collect_headings(sub, num, indent_level + 1)
                    entries.extend(sub_entries)
        return entries

    toc_entries = _collect_headings(sections)

    for entry in toc_entries:
        level = entry['level']
        indent = '  ' * (level - 1)
        text = f"{indent}{entry['number']} {entry['title']}"

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 1.3)

        if level == 1:
            run = p.add_run(text)
            set_run_font(run, HEI_TI, TIMES, XIAO_SI, bold=True)
        else:
            run = p.add_run(text)
            set_run_font(run, SONG_TI, TIMES, XIAO_SI)

        tab_run = p.add_run('\t')
        set_run_font(tab_run, SONG_TI, TIMES, XIAO_SI)

    extras = []
    refs = data.get('references', [])
    appendices = data.get('appendices', [])
    ack = data.get('acknowledgement', '')

    if refs:
        extras.append(('参考文献', 1))
    if appendices:
        for i, app in enumerate(appendices):
            extras.append((app.get('title', f'附录{i+1}'), 2))
    if ack:
        extras.append(('致谢', 1))

    for title, level in extras:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 1.3)
        run = p.add_run(title)
        set_run_font(run, HEI_TI if level == 1 else SONG_TI, TIMES, XIAO_SI,
                      bold=(level == 1))

    doc.add_page_break()


def generate_body(doc, data):
    """Generate body sections."""
    sections = data.get('sections', [])
    section_counter = [0, 0, 0]

    def process_children(children, depth=0):
        for item in children:
            if item.get('type') == 'section' or ('level' in item and 'title' in item):
                level = item.get('level', 1)

                if level == 1:
                    section_counter[0] += 1
                    section_counter[1] = 0
                    section_counter[2] = 0
                    num = str(section_counter[0])
                elif level == 2:
                    section_counter[1] += 1
                    section_counter[2] = 0
                    num = f"{section_counter[0]}.{section_counter[1]}"
                else:
                    section_counter[2] += 1
                    num = f"{section_counter[0]}.{section_counter[1]}.{section_counter[2]}"

                add_section_heading(doc, num, item.get('title', ''), level)

                sub_children = item.get('children', [])
                if sub_children:
                    process_children(sub_children, depth + 1)

            elif item.get('type') == 'paragraph':
                text = item.get('text', '')
                p = doc.add_paragraph()
                set_paragraph_spacing(p, 1.5)
                _add_text_with_refs(p, text)

            elif item.get('type') == 'table':
                _add_table(doc, item)

            elif item.get('type') == 'figure':
                _add_figure(doc, item)

            elif item.get('type') == 'equation':
                _add_equation(doc, item)

    process_children(sections)


def _add_text_with_refs(paragraph, text):
    """Add text with superscript references like [1], [2,3], [1-3]."""
    import re
    pattern = r'(\[\d+(?:[-,]\d+)*\])'
    parts = re.split(pattern, text)

    for part in parts:
        if re.match(pattern, part):
            run = paragraph.add_run(part)
            set_run_font(run, SONG_TI, TIMES, XIAO_SI)
            run.font.superscript = True
        else:
            run = paragraph.add_run(part)
            set_run_font(run, SONG_TI, TIMES, XIAO_SI)


def _add_table(doc, item):
    """Add a 三线表 (three-line table)."""
    caption = item.get('caption', '表')
    headers = item.get('headers', [])
    rows = item.get('rows', [])
    notes = item.get('notes', [])

    p_space = doc.add_paragraph()
    set_paragraph_spacing(p_space, 1.5, space_before=6, space_after=3)

    p_cap = add_paragraph_with_font(
        doc, caption, SONG_TI, TIMES, XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    p_cap.paragraph_format.space_after = Pt(3)

    n_rows = len(rows) + 1
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for j, header in enumerate(headers):
        format_table_cell(table.rows[0].cells[j], header, bold=True)

    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            format_table_cell(table.rows[i + 1].cells[j], cell_data)

    apply_three_line_table(table)

    if notes:
        for note in notes:
            p_note = doc.add_paragraph()
            set_paragraph_spacing(p_note, 1.5)
            run_label = p_note.add_run('注：')
            set_run_font(run_label, SONG_TI, TIMES, WU_HAO, bold=True)
            run_text = p_note.add_run(note)
            set_run_font(run_text, SONG_TI, TIMES, WU_HAO)


def _add_figure(doc, item):
    """Add a figure caption (placeholder)."""
    caption = item.get('caption', '图')
    p = add_paragraph_with_font(
        doc, caption, SONG_TI, TIMES, XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    p_note = add_paragraph_with_font(
        doc, '[图片插入位置]', SONG_TI, TIMES, XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    set_paragraph_spacing(p_note, 1.5, space_before=3, space_after=6)


def _add_equation(doc, item):
    """Add an equation (centered with right-aligned number)."""
    latex = item.get('latex', '')
    eq_number = item.get('number', '')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 1.5, space_before=3, space_after=3)

    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(7.935))
    pf.tab_stops.add_tab_stop(Cm(15.87), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    run_tab1 = p.add_run('\t')
    set_run_font(run_tab1, TIMES, TIMES, XIAO_SI)
    run_eq = p.add_run(f'  {latex}  ')
    set_run_font(run_eq, TIMES, TIMES, XIAO_SI, italic=True)
    if eq_number:
        run_tab2 = p.add_run('\t')
        set_run_font(run_tab2, TIMES, TIMES, XIAO_SI)
        run_num = p.add_run(f'({eq_number})')
        set_run_font(run_num, TIMES, TIMES, XIAO_SI)


def generate_references(doc, data):
    """Generate references section."""
    refs = data.get('references', [])
    if not refs:
        return

    doc.add_page_break()

    p_ref_title = add_paragraph_with_font(
        doc, '参考文献', HEI_TI, TIMES, WU_HAO,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    p_ref_title.paragraph_format.space_after = Pt(6)

    for i, ref in enumerate(refs, 1):
        ref_text = build_reference_text(ref, i)
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 1.5)
        run = p.add_run(ref_text)
        set_run_font(run, SONG_TI, TIMES, WU_HAO)


def generate_appendices(doc, data):
    """Generate appendices."""
    appendices = data.get('appendices', [])
    for app in appendices:
        doc.add_page_break()
        title = app.get('title', '附录')
        p = add_paragraph_with_font(
            doc, title, HEI_TI, TIMES, SI_HAO,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT
        )
        p.paragraph_format.space_after = Pt(6)

        for content_item in app.get('content', []):
            p2 = doc.add_paragraph()
            set_paragraph_spacing(p2, 1.5)
            run = p2.add_run(content_item)
            set_run_font(run, SONG_TI, TIMES, XIAO_SI)


def generate_acknowledgement(doc, data):
    """Generate acknowledgement section."""
    ack = data.get('acknowledgement', '')
    if not ack:
        return

    doc.add_page_break()
    p_ack_title = add_paragraph_with_font(
        doc, '致谢', HEI_TI, TIMES, SI_HAO,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    p_ack_title.paragraph_format.space_after = Pt(6)

    p_ack = doc.add_paragraph()
    set_paragraph_spacing(p_ack, 1.5)
    run = p_ack.add_run(ack)
    set_run_font(run, SONG_TI, TIMES, XIAO_SI)


def generate_paper(input_file, output_file):
    """Generate a formatted Chinese academic paper .docx from JSON input."""
    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = TIMES
    style.font.size = XIAO_SI
    style.element.rPr.rFonts.set(qn('w:eastAsia'), SONG_TI)

    generate_front_matter(doc, data)

    new_section = doc.add_section()
    new_section.top_margin = Cm(2.54)
    new_section.bottom_margin = Cm(2.54)
    new_section.left_margin = Cm(3.17)
    new_section.right_margin = Cm(3.17)
    sectPr = new_section._sectPr
    pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="decimal"/>')
    sectPr.append(pgNumType)

    footer = new_section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in [run, run2, run3]:
        set_run_font(r, SONG_TI, TIMES, LIU_HAO)

    generate_toc(doc, data)
    generate_body(doc, data)
    generate_references(doc, data)
    generate_appendices(doc, data)
    generate_acknowledgement(doc, data)

    doc.save(output_file)
    print(f"Paper generated successfully: {output_file}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python generate_paper.py input.json output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    generate_paper(input_file, output_file)
